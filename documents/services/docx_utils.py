# documents/services/docx_utils.py
# Общая логика для генерации договоров через {{ключ}}-теги с сохранением
# исходного форматирования шаблона (шрифт/размер/жирность берутся с самого
# документа, а не хардкодятся). Используется всеми типами договоров.

from docx.oxml.ns import qn


def apply_font_from(run, source_run):
    """Копирует шрифт/размер/жирность с исходного run, сохраняя оформление шаблона."""
    font_name = source_run.font.name or "Times New Roman"
    font_size = source_run.font.size

    run.font.name = font_name
    if font_size:
        run.font.size = font_size
    run.font.bold = source_run.font.bold
    run.font.italic = source_run.font.italic
    run.font.underline = source_run.font.underline

    rFonts = run._element.rPr.rFonts
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)


def replace_text_preserving_format(doc, data):
    """Заменяет плейсхолдеры вида {{ключ}}, сохраняя исходное форматирование абзаца."""

    def process_paragraph(paragraph):
        full_text = paragraph.text
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if value is None:
                value = ""
            full_text = full_text.replace(placeholder, str(value))

        if full_text == paragraph.text:
            return

        source_run = paragraph.runs[0] if paragraph.runs else None

        for _ in range(len(paragraph.runs)):
            paragraph.runs[0].clear()
            paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)

        run = paragraph.add_run(full_text)
        if source_run is not None:
            apply_font_from(run, source_run)

    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
