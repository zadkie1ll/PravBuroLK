# documents/services/mfc_contract.py

import base64
import os
from datetime import datetime

import requests
from docx import Document
from docx.oxml.ns import qn

from documents.views import (
    WEBHOOK_URL,
    insert_table_after_heading,
    number_to_words,
)

MFC_CONTRACT_FIELD = "UF_CRM_1787661784627"


def _apply_font_from(run, source_run):
    """Копирует шрифт/размер/жирность с исходного run, сохраняя оформление шаблона."""
    font_name = source_run.font.name or "Times New Roman"
    font_size = source_run.font.size

    run.font.name = font_name
    if font_size:
        run.font.size = font_size
    run.font.bold = source_run.font.bold
    run.font.italic = source_run.font.italic
    run.font.underline = source_run.font.underline

    rFonts = run._element.rPr.rFonts
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)


def replace_text_in_paragraphs_mfc(doc, data):
    """Заменяет плейсхолдеры вида {{ключ}}, сохраняя исходное форматирование абзаца."""

    def process_paragraph(paragraph):
        full_text = paragraph.text
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if value is None:
                value = ""
            full_text = full_text.replace(placeholder, str(value))

        if full_text == paragraph.text:
            return

        source_run = paragraph.runs[0] if paragraph.runs else None

        for _ in range(len(paragraph.runs)):
            paragraph.runs[0].clear()
            paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)

        run = paragraph.add_run(full_text)
        if source_run is not None:
            _apply_font_from(run, source_run)

    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)


def generate_mfc_contract(data, template_path, output_path, payments):
    today = datetime.now().strftime("%d.%m.%Y")
    data["today"] = today
    discount = int(data.get("скидка", 0))

    if data.get("сумма бонус", "") == "":
        data["сумма бонус"] = "0"

    fio_parts = data.get("ФИО", "").split()
    if len(fio_parts) < 3:
        raise ValueError(f"ФИО должно содержать минимум 3 части (Фамилия Имя Отчество), получили: {data.get('ФИО')}")

    data["инициалы"] = f"{fio_parts[0]} {fio_parts[1][0]}. {fio_parts[2][0]}."

    data["сумма юристы"] = str(int(data["сумма юристы"]) - discount)
    data["сумма"] = str(int(data["сумма бонус"]) + int(data["сумма юристы"]))
    data["words_sum"] = number_to_words(int(data["сумма юристы"]))

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)
    replace_text_in_paragraphs_mfc(doc, data)
    insert_table_after_heading(doc, payments)

    doc.save(output_path)


def upload_mfc_contract_to_bitrix(deal_id, file_path, field_id):
    """Заливает МФЦ-договор в выделенное под него поле сделки Bitrix (UF_CRM_1787661784627)."""
    if not os.path.exists(file_path):
        return {"status": "error", "message": f"Файл не найден: {file_path}"}

    with open(file_path, "rb") as f:
        encoded = base64.b64encode(f.read()).decode("utf-8")

    payload = {
        "id": deal_id,
        "fields": {field_id: {"fileData": [os.path.basename(file_path), encoded]}},
    }

    url = f"{WEBHOOK_URL.rstrip('/')}/crm.deal.update.json"
    response = requests.post(url, json=payload)

    try:
        result = response.json()
    except Exception:
        return {"status": "error", "message": "Bitrix вернул не-JSON", "raw": response.text}

    if result.get("error"):
        return {
            "status": "error",
            "message": result.get("error_description", "Ошибка Bitrix"),
            "details": result,
        }

    return {"status": "success", "message": "Файл загружен"}
