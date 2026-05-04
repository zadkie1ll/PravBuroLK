import base64
import json
import logging
import mimetypes
import os
import re
import time
import traceback
from copy import deepcopy
from datetime import date, datetime
from decimal import Decimal
from pathlib import Path
from urllib.parse import quote, urlencode, urlsplit

import requests
import telebot
from dateutil.relativedelta import relativedelta
from django.conf import settings
from django.core import signing
from django.db.models import Sum
from django.forms.models import model_to_dict
from django.http import FileResponse, Http404, HttpResponse, JsonResponse
from django.shortcuts import redirect, render
from django.urls import reverse
from django.utils.timezone import now
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from .services.document_pipeline import DocumentPipeline

USE_PROXIES = True

WEBHOOK_URL = "https://prav-buro.bitrix24.ru/rest/24/vszzr53045oedn5m/"
CONTRACT_FILE_FIELD = "UF_CRM_1745892619372"
CONTRACT_ACCEPTED_FIELD = "UF_CRM_1775216196958"
CONTRACT_LINK_FIELD = "UF_CRM_1775217002"
CONTRACT_NUMBER_FIELD = "UF_CRM_1745892727271"
CONTRACT_FIRST_PAYMENT_FIELD = "UF_CRM_1742468532579"
CONTRACT_PAGE_SIGN_SALT = "documents.contract.confirmation"

logger = logging.getLogger(__name__)


def _iter_bitrix_webhook_urls(*urls: str | None):
    seen = set()
    for url in urls:
        if not url:
            continue
        normalized = url.rstrip("/") + "/"
        if normalized in seen:
            continue
        seen.add(normalized)
        yield normalized


def _iter_bitrix_portal_urls() -> list[str]:
    portals = []
    seen = set()
    for webhook_url in _iter_bitrix_webhook_urls(
        WEBHOOK_URL,
        getattr(settings, "BITRIX_WEBHOOK_URL", None),
    ):
        parts = urlsplit(webhook_url)
        if not parts.scheme or not parts.netloc:
            continue
        portal_url = f"{parts.scheme}://{parts.netloc}"
        if portal_url in seen:
            continue
        seen.add(portal_url)
        portals.append(portal_url)
    return portals


def generate_document(request):
    if request.method != "POST":
        raise Http404("Этот эндпоинт принимает только POST-запросы.")

    context = request.POST.dict()

    # --- КРЕДИТОРЫ ---
    creditors_data, total_debt, distribution_sum = parse_creditors_and_calculate(request)
    context["debts"] = creditors_data
    context["total_debt"] = total_debt
    context["total_pay"] = distribution_sum     

    # --- ВОЗМЕЩЕНИЯ УБЫТКОВ ---
    losses_names = request.POST.getlist("loss_org[]")
    losses_amounts = request.POST.getlist("loss_amount[]")

    losses = []
    total_losses = 0

    for name, amount in zip(losses_names, losses_amounts):
        if name.strip() or amount.strip():
            amount_float = float(amount) if amount.strip() else 0
            losses.append({
                "name": name,
                "amount": amount_float
            })
            total_losses += amount_float

    context["losses"] = losses
    context["total_losses"] = total_losses

    context["grand_total"] = total_debt + total_losses

    # ----ПАЙПЛАЙН------------------------------------------------------------------------------------------------------------------------
    template_path = os.path.join(
        settings.BASE_DIR,
        "documents",
        "templates_src",
        "test_template.docx"
    )

    output_dir = os.path.join(settings.MEDIA_ROOT, "generated_docs")
    os.makedirs(output_dir, exist_ok=True)

    output_path = os.path.join(output_dir, "generated_document.docx")

    pipeline = DocumentPipeline(template_path, context)
    pipeline.run()
    pipeline.save(output_path)

    with open(output_path, "rb") as file:
        response = HttpResponse(
            file.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response['Content-Disposition'] = 'attachment; filename="generated_document.docx"'
        return response


def document_form(request):
    """
    Просто рендерит страницу — форма должна отправлять в generate_document.
    """
    return render(request, "document_form.html")


def _get_documents_dir() -> Path:
    return Path(settings.BASE_DIR) / "documents"


def _get_generated_contract_path(deal_id: str | int) -> Path:
    return _get_documents_dir() / "generated_docs" / f"dogovor_{deal_id}.docx"


def _extract_deal_id(post_data) -> str | None:
    document_id_2 = post_data.get("document_id[2]")
    if not document_id_2:
        return None

    match = re.search(r"DEAL_(\d+)", document_id_2)
    if not match:
        return None

    return match.group(1)


def _bitrix_get(
    method: str,
    params: dict | None = None,
    timeout: int = 10,
    webhook_url: str | None = None,
) -> dict:
    response = requests.get(
        f"{(webhook_url or WEBHOOK_URL).rstrip('/')}/{method}.json",
        params=params or {},
        timeout=timeout,
    )
    response.raise_for_status()
    payload = response.json()
    if payload.get("error"):
        raise RuntimeError(payload.get("error_description") or payload["error"])
    return payload


def _bitrix_post(method: str, payload: dict, timeout: int = 10) -> dict:
    response = requests.post(
        f"{WEBHOOK_URL.rstrip('/')}/{method}.json",
        json=payload,
        timeout=timeout,
    )
    response.raise_for_status()
    data = response.json()
    if data.get("error"):
        raise RuntimeError(data.get("error_description") or data["error"])
    return data


def _get_deal_data(deal_id: str) -> dict:
    payload = _bitrix_get("crm.deal.get", {"ID": deal_id})
    deal_data = payload.get("result")
    if not deal_data:
        raise RuntimeError("Deal not found")
    return deal_data


def _build_contract_token(deal_id: str) -> str:
    signer = signing.Signer(salt=CONTRACT_PAGE_SIGN_SALT)
    return signer.sign(deal_id)


def _validate_contract_token(deal_id: str, token: str | None) -> bool:
    if not token:
        return False

    signer = signing.Signer(salt=CONTRACT_PAGE_SIGN_SALT)
    try:
        return signer.unsign(token) == str(deal_id)
    except signing.BadSignature:
        return False


def _extract_file_id(file_value):
    if not file_value:
        return None

    if isinstance(file_value, dict):
        for key in ("id", "ID", "fileId", "FILE_ID"):
            if file_value.get(key):
                return _extract_file_id(file_value.get(key))
        return None

    if isinstance(file_value, (list, tuple)):
        for item in file_value:
            file_id = _extract_file_id(item)
            if file_id is not None:
                return file_id
        return None

    if isinstance(file_value, str):
        if file_value.startswith(("http://", "https://")):
            return file_value
        if file_value.isdigit():
            return int(file_value)
        return None

    if isinstance(file_value, int):
        return file_value

    return None


def _extract_file_url(file_value) -> str | None:
    if not file_value:
        return None

    if isinstance(file_value, dict):
        for key in (
            "url",
            "URL",
            "downloadUrl",
            "DOWNLOAD_URL",
            "showUrl",
            "SHOW_URL",
            "detailUrl",
            "DETAIL_URL",
            "src",
            "SRC",
        ):
            value = file_value.get(key)
            if isinstance(value, str):
                if value.startswith(("http://", "https://")):
                    return value
                if value.startswith("/"):
                    portal_urls = _iter_bitrix_portal_urls()
                    if portal_urls:
                        return f"{portal_urls[0]}{value}"

        for value in file_value.values():
            nested_url = _extract_file_url(value)
            if nested_url:
                return nested_url
        return None

    if isinstance(file_value, (list, tuple)):
        for item in file_value:
            nested_url = _extract_file_url(item)
            if nested_url:
                return nested_url
        return None

    if isinstance(file_value, str) and file_value.startswith(("http://", "https://")):
        return file_value

    return None


def _resolve_contract_download_url(file_value) -> str | None:
    direct_url = _extract_file_url(file_value)
    if direct_url:
        return direct_url

    file_id = _extract_file_id(file_value)
    if not file_id:
        return None

    if isinstance(file_id, str) and file_id.startswith(("http://", "https://")):
        return file_id

    last_error = None
    webhook_candidates = list(
        _iter_bitrix_webhook_urls(
            WEBHOOK_URL,
            getattr(settings, "BITRIX_WEBHOOK_URL", None),
        )
    )

    for webhook_url in webhook_candidates:
        try:
            payload = _bitrix_get(
                "disk.file.get",
                {"id": file_id},
                webhook_url=webhook_url,
            )
            file_data = payload.get("result") or {}
            resolved_url = (
                file_data.get("DOWNLOAD_URL")
                or file_data.get("DETAIL_URL")
                or file_data.get("SRC")
            )
            if resolved_url:
                return resolved_url
        except Exception as exc:
            last_error = exc
            logger.warning(
                "Failed to resolve Bitrix disk file URL for file_id=%s via webhook=%s: %s",
                file_id,
                webhook_url,
                exc,
            )

    if last_error:
        logger.error(
            "All Bitrix webhook attempts failed for file_id=%s: %s",
            file_id,
            last_error,
        )

    return None


def _build_contract_page_url(request, deal_id: str) -> str:
    token = _build_contract_token(deal_id)
    base_url = request.build_absolute_uri(reverse("contract_confirmation_page", args=[deal_id]))
    return f"{base_url}?token={quote(token)}"


def _build_contract_file_url(request, deal_id: str | int, token: str) -> str:
    base_url = request.build_absolute_uri(reverse("contract_document_file", args=[deal_id]))
    return f"{base_url}?token={quote(token)}"


def _build_contract_payment_url(request, deal_id: str | int, token: str) -> str:
    base_url = request.build_absolute_uri(reverse("contract_payment_redirect", args=[deal_id]))
    return f"{base_url}?token={quote(token)}"


def _build_office_preview_url(download_url: str | None) -> str | None:
    if not download_url:
        return None
    return f"https://view.officeapps.live.com/op/embed.aspx?src={quote(download_url, safe='')}"


def _extract_decimal_amount(value) -> Decimal:
    if value in (None, "", False):
        return Decimal("0")
    if isinstance(value, (int, float, Decimal)):
        return Decimal(str(value))

    normalized = str(value).split("|")[0].strip().replace(" ", "").replace(",", ".")
    if not normalized:
        return Decimal("0")
    return Decimal(normalized)


def _format_amount_rub(amount: Decimal) -> str:
    normalized = amount.quantize(Decimal("1")) if amount == amount.to_integral_value() else amount.quantize(Decimal("0.01"))
    return f"{normalized:,.2f}".replace(",", " ").replace(".00", "")


def _get_contract_payment_purpose(deal_data: dict) -> str:
    client_name = (deal_data.get("TITLE") or "").strip()
    contract_number = (deal_data.get(CONTRACT_NUMBER_FIELD) or "").strip()
    if contract_number and client_name:
        return f"Оплата по договору {contract_number} {client_name}"
    if contract_number:
        return f"Оплата по договору {contract_number}"
    return f"Оплата по договору {client_name}".strip()


def _get_alfa_register_url() -> str:
    configured = getattr(settings, "ALFA_API_URL_PROD", "").strip()
    if configured:
        return configured if configured.endswith(".do") else f"{configured.rstrip('/')}/register.do"
    return "https://payment.alfabank.ru/payment/rest/register.do"


def _get_alfa_username() -> str:
    return getattr(settings, "ALFA_USER_PROD", "").strip() or "r-prav_0-api"


def _get_alfa_password() -> str:
    return getattr(settings, "ALFA_PASS_PROD", "").strip() or "Qwasdcvbgh243567!@"


def _build_contract_page_return_url(request, deal_id: str | int, token: str, payment_state: str | None = None) -> str:
    base_url = request.build_absolute_uri(reverse("contract_confirmation_page", args=[deal_id]))
    query = {"token": token}
    if payment_state:
        query["payment_state"] = payment_state
    return f"{base_url}?{urlencode(query)}"


def _register_contract_payment(request, deal_id: str, token: str, deal_data: dict) -> str:
    amount = _extract_decimal_amount(deal_data.get(CONTRACT_FIRST_PAYMENT_FIELD))
    if amount <= 0:
        raise RuntimeError("Сумма первого платежа не заполнена")
    order_number = f"contract-{deal_id}-{int(time.time())}"

    payload = {
        "userName": _get_alfa_username(),
        "password": _get_alfa_password(),
        "orderNumber": order_number,
        "amount": int((amount * 100).quantize(Decimal("1"))),
        "description": _get_contract_payment_purpose(deal_data),
        "returnUrl": f"{_build_contract_page_return_url(request, deal_id, token, 'success')}&orderNumber={quote(order_number)}",
        "failUrl": f"{_build_contract_page_return_url(request, deal_id, token, 'failed')}&orderNumber={quote(order_number)}",
    }

    response = requests.post(_get_alfa_register_url(), data=payload, timeout=15)
    response.raise_for_status()
    data = response.json()

    if data.get("errorCode") and str(data["errorCode"]) != "0":
        raise RuntimeError(data.get("errorMessage") or "Не удалось создать оплату в Альфа-Банке")

    form_url = data.get("formUrl")
    if not form_url:
        raise RuntimeError("Альфа-Банк не вернул ссылку на оплату")
    return form_url


def _get_contract_payment_qr_data_uri() -> str:
    image_path = _get_documents_dir() / "2026-04-03 16.29.44.jpg"
    if not image_path.exists():
        return ""

    mime_type, _ = mimetypes.guess_type(image_path.name)
    encoded = base64.b64encode(image_path.read_bytes()).decode("ascii")
    return f"data:{mime_type or 'image/jpeg'};base64,{encoded}"


def _get_contract_payment_requisites() -> list[tuple[str, str]]:
    requisites = [
        ("Получатель", getattr(settings, "CONTRACT_PAYMENT_RECIPIENT", "").strip()),
        ("Адрес", getattr(settings, "CONTRACT_PAYMENT_ADDRESS", "").strip()),
        ("ИНН", getattr(settings, "CONTRACT_PAYMENT_INN", "").strip()),
        ("КПП", getattr(settings, "CONTRACT_PAYMENT_KPP", "").strip()),
        ("Валюта", getattr(settings, "CONTRACT_PAYMENT_CURRENCY", "").strip()),
        ("Банк", getattr(settings, "CONTRACT_PAYMENT_BANK", "").strip()),
        ("БИК", getattr(settings, "CONTRACT_PAYMENT_BIK", "").strip()),
        ("Расчетный счет", getattr(settings, "CONTRACT_PAYMENT_ACCOUNT", "").strip()),
        ("Корреспондентский счет", getattr(settings, "CONTRACT_PAYMENT_CORR_ACCOUNT", "").strip()),
    ]
    return [(label, value) for label, value in requisites if value]


def _build_contract_payment_context(request, deal_id: str | int, token: str, deal_data: dict) -> dict:
    first_payment_amount = _extract_decimal_amount(deal_data.get(CONTRACT_FIRST_PAYMENT_FIELD))
    return {
        "contract_payment_url": _build_contract_payment_url(request, deal_id, token),
        "first_payment_amount": _format_amount_rub(first_payment_amount),
        "payment_purpose": _get_contract_payment_purpose(deal_data),
        "payment_requisites": _get_contract_payment_requisites(),
        "payment_qr_data_uri": _get_contract_payment_qr_data_uri(),
    }


def _add_contract_payment_timeline_comment(deal_id: str | int, order_number: str, amount: Decimal) -> None:
    response = requests.post(
        f"{WEBHOOK_URL.rstrip('/')}/crm.timeline.comment.add",
        data={
            "fields[ENTITY_ID]": str(deal_id),
            "fields[ENTITY_TYPE]": "deal",
            "fields[COMMENT]": (
                "Поступила успешная оплата по договору\n"
                f"Сумма: {_format_amount_rub(amount)} ₽\n"
                f"Номер платежа: {order_number}"
            ),
        },
        timeout=15,
    )
    response.raise_for_status()
    payload = response.json()
    if payload.get("error"):
        raise RuntimeError(payload.get("error_description") or payload["error"])


def _update_contract_confirmation(deal_id: str, accepted: bool = True) -> dict:
    return _bitrix_post(
        "crm.deal.update",
        {
            "id": deal_id,
            "fields": {
                CONTRACT_ACCEPTED_FIELD: 1 if accepted else 0,
            },
        },
    )


def _update_contract_link(deal_id: str, contract_url: str) -> dict:
    return _bitrix_post(
        "crm.deal.update",
        {
            "id": deal_id,
            "fields": {
                CONTRACT_LINK_FIELD: contract_url,
            },
        },
    )


def contract_confirmation_page(request, deal_id: int):
    token = request.GET.get("token") or request.POST.get("token")
    if not _validate_contract_token(str(deal_id), token):
        raise Http404("Страница договора не найдена")

    error_message = ""
    success_message = ""

    try:
        deal_data = _get_deal_data(str(deal_id))
    except Exception as exc:
        logger.exception("Failed to load contract page for deal %s", deal_id)
        return render(
            request,
            "documents/contract_confirmation.html",
            {
                "deal_id": deal_id,
                "token": token,
                "error_message": f"Не удалось загрузить договор: {exc}",
                "contract_download_url": None,
                "contract_preview_url": None,
                "is_confirmed": False,
                "contract_number": "",
                "client_name": "",
            },
            status=502,
        )

    if request.method == "POST":
        if request.POST.get("agree") != "on":
            error_message = "Для подтверждения нужно отметить согласие с договором."
        else:
            try:
                _update_contract_confirmation(str(deal_id), accepted=True)
                success_message = "Договор подтвержден."
                deal_data[CONTRACT_ACCEPTED_FIELD] = 1
            except Exception as exc:
                logger.exception("Failed to confirm contract for deal %s", deal_id)
                error_message = f"Не удалось сохранить подтверждение: {exc}"

    contract_download_url = None
    generated_contract_path = _get_generated_contract_path(deal_id)
    if generated_contract_path.exists():
        contract_download_url = _build_contract_file_url(request, deal_id, token)
    else:
        contract_download_url = _resolve_contract_download_url(deal_data.get(CONTRACT_FILE_FIELD))
    if not contract_download_url and not error_message:
        error_message = "Не удалось получить прямую ссылку на файл договора. Подтверждение остается доступным."
    contract_preview_url = _build_office_preview_url(contract_download_url)
    is_confirmed = str(deal_data.get(CONTRACT_ACCEPTED_FIELD)).upper() in {"1", "Y", "TRUE"}
    payment_state = (request.GET.get("payment_state") or "").strip().lower()
    payment_error = (request.GET.get("payment_error") or "").strip()
    payment_order_number = (request.GET.get("orderNumber") or "").strip()
    session_payment_key = f"contract_payment_success_{deal_id}"
    is_payment_completed = bool(request.session.get(session_payment_key))

    if payment_state == "success":
        is_payment_completed = True
        request.session[session_payment_key] = True

        session_log_key = f"contract_payment_logged_{payment_order_number or deal_id}"
        if payment_order_number and not request.session.get(session_log_key):
            try:
                _add_contract_payment_timeline_comment(
                    deal_id=deal_id,
                    order_number=payment_order_number,
                    amount=_extract_decimal_amount(deal_data.get(CONTRACT_FIRST_PAYMENT_FIELD)),
                )
                request.session[session_log_key] = True
            except Exception as exc:
                logger.exception("Failed to add contract payment timeline comment for deal %s", deal_id)
                if not payment_error:
                    payment_error = f"Не удалось записать оплату в журнал: {exc}"

    payment_context = _build_contract_payment_context(request, deal_id, token, deal_data)

    return render(
        request,
        "documents/contract_confirmation.html",
        {
            "deal_id": deal_id,
            "token": token,
            "error_message": error_message,
            "success_message": success_message,
            "contract_download_url": contract_download_url,
            "contract_preview_url": contract_preview_url,
            "is_confirmed": is_confirmed,
            "contract_number": deal_data.get(CONTRACT_NUMBER_FIELD, ""),
            "client_name": deal_data.get("TITLE", ""),
            "payment_state": payment_state,
            "payment_error": payment_error,
            "is_payment_completed": is_payment_completed,
            **payment_context,
        },
    )


def contract_document_file(request, deal_id: int):
    token = request.GET.get("token")
    if not _validate_contract_token(str(deal_id), token):
        raise Http404("Файл договора не найден")

    contract_path = _get_generated_contract_path(deal_id)
    if not contract_path.exists():
        raise Http404("Файл договора не найден")

    return FileResponse(
        contract_path.open("rb"),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=False,
        filename=contract_path.name,
    )


def contract_payment_redirect(request, deal_id: int):
    token = request.GET.get("token")
    if not _validate_contract_token(str(deal_id), token):
        raise Http404("Страница оплаты не найдена")

    try:
        deal_data = _get_deal_data(str(deal_id))
        form_url = _register_contract_payment(request, str(deal_id), token, deal_data)
    except Exception as exc:
        logger.exception("Failed to create contract payment for deal %s", deal_id)
        error_query = urlencode(
            {
                "token": token,
                "payment_state": "error",
                "payment_error": str(exc),
            }
        )
        return redirect(f"{reverse('contract_confirmation_page', args=[deal_id])}?{error_query}")

    return redirect(form_url)




@csrf_exempt
def dogovor(request):
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'message': 'Only POST allowed'}, status=405)

    logger.error(f"INCOMING POST: {request.POST}")

    # --- Получение ID сделки ---
    deal_id = _extract_deal_id(request.POST)
    if not deal_id:
        return JsonResponse({'status': 'error', 'message': 'Invalid deal ID'}, status=400)

    try:
        deal_data = _get_deal_data(deal_id)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'Deal fetch error: {e}'}, status=500)

    try:
        contact_id = deal_data.get("CONTACT_ID")
        phone_number = get_phone_number(contact_id) if contact_id else "00000000000"

        fio = deal_data.get("TITLE", "")
        fio_parts = fio.split()

        last_name  = fio_parts[0] if len(fio_parts) > 0 else ""
        first_name = fio_parts[1] if len(fio_parts) > 1 else ""
        mid_name   = fio_parts[2] if len(fio_parts) > 2 else ""

        contract = {
            "номер договора": deal_data.get('UF_CRM_1745892727271', "000"),
            "ФИО": fio,
            "фамилия": last_name,
            "имя": first_name,
            "отчество": mid_name,
            "дата рождения": format_date(deal_data.get('UF_CRM_1745888327609')),
            "серия": deal_data.get('UF_CRM_1745889060779', ''),
            "номер": deal_data.get('UF_CRM_1745889067225', ''),
            "кем": deal_data.get('UF_CRM_1745889085935', ''),
            "дата выдачи": format_date(deal_data.get('UF_CRM_1754384630146')),
            "код": deal_data.get('UF_CRM_1745889094660', ''),
            "место рождения": deal_data.get('UF_CRM_1745889105838', ''),
            "адрес регистрации": deal_data.get('UF_CRM_1745893079148', ''),
            "номер телефона": phone_number,
            "сумма юристы": str(int(float(deal_data.get('OPPORTUNITY', 0)))),
            "сумма бонус": deal_data.get('UF_CRM_1742457114242', "0").split("|")[0],
            "Первый платеж": deal_data.get('UF_CRM_1742468532579', "0").split("|")[0],
            "today": datetime.today().strftime("%d.%m.%Y"),
            "data": datetime.today().strftime("%m/%Y"),
            "количество платежей": deal_data.get('UF_CRM_1742480133860', "1"),
            "скидка": deal_data.get('UF_CRM_1742457148727', "0").split("|")[0],
            "дата начала платежей": format_date(deal_data.get('UF_CRM_1742468566169')),
            "Число для оплаты": deal_data.get('UF_CRM_1745893194511', "1"),
        }

    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'Data preparation error: {e}'}, status=500)

    try:
        total = int(contract["сумма юристы"]) + int(contract["сумма бонус"])
        num = int(contract["количество платежей"])
        first = int(contract["Первый платеж"])
        discount = int(contract["скидка"])
        start_date = contract["дата начала платежей"]
        second_day = contract["Число для оплаты"]

        payments = calculate_payments(
            num, total, discount, start_date, first,
            second_payment_day=second_day
        )

    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'Payment calc error: {e}'}, status=500)

    # --- KРОССПЛАТФОРМЕННЫЕ ПУТИ ---
    BASE = _get_documents_dir()

    template = BASE / "templates_src" / "template_2.docx"
    output = BASE / "generated_docs" / f"dogovor_{deal_id}.docx"

    logger.error(f"TEMPLATE PATH: {template}")
    logger.error(f"OUTPUT PATH: {output}")

    # Проверка шаблона
    if not template.exists():
        return JsonResponse({'status': 'error', 'message': f'Template not found: {template}'}, status=500)

    # --- Генерация договора ---
    try:
        output.parent.mkdir(parents=True, exist_ok=True)

        generate_contract(contract, str(template), str(output), payments)

        logger.error(f"GENERATED FILE EXISTS: {output.exists()}")
        if output.exists():
            logger.error(f"FILE SIZE: {output.stat().st_size} bytes")

    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'Doc generation error: {e}'}, status=500)
    try:
        result = upload_to_bitrix(
            deal_id,
            str(output),
            CONTRACT_FILE_FIELD,
            payments
        )
        logger.error(f"BITRIX UPLOAD RESULT: {result}")

    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'Upload error: {e}'}, status=500)

    confirmation_url = _build_contract_page_url(request, deal_id)
    contract_token = _build_contract_token(deal_id)
    contract_download_url = _build_contract_file_url(request, deal_id, contract_token)
    contract_preview_url = _build_office_preview_url(contract_download_url)

    try:
        _update_contract_link(deal_id, confirmation_url)
    except Exception:
        logger.exception("Failed to save or resolve contract URL for deal %s", deal_id)

    return JsonResponse(
        {
            'status': 'success',
            'message': 'Document generated & uploaded',
            'deal_id': deal_id,
            'confirmation_url': confirmation_url,
            'contract_download_url': contract_download_url,
            'contract_preview_url': contract_preview_url,
        }
    )

def calculate_payments(num_payments, total_amount, discount, start_date, first_payment, second_payment_day):
    if num_payments == 1 and first_payment >= (total_amount - discount):
        return [[1, start_date, f"{first_payment:.2f}"]]

    remaining_amount = (total_amount - discount) - first_payment
    remaining_payments = num_payments - 1

    if remaining_payments == 0:
        return [[1, start_date, f"{first_payment:.2f}"]]

    payment_amount = round(remaining_amount / remaining_payments, -2)
    
    adjusted_total = first_payment + payment_amount * (remaining_payments - 1)
    last_payment = (total_amount - discount) - adjusted_total 

    first_date = datetime.strptime(start_date, "%d.%m.%Y").date()
    table_data = [[1, first_date.strftime("%d.%m.%Y"), f"{first_payment:.2f}"]]

    second_payment_day = int(second_payment_day)  
    second_date = (first_date + relativedelta(months=1)).replace(day=second_payment_day)

    while second_date.day != second_payment_day:
        second_date -= relativedelta(days=1)

    table_data.append([2, second_date.strftime("%d.%m.%Y"), f"{payment_amount:.2f}"])

    current_date = second_date
    for i in range(2, num_payments - 1):  
        current_date += relativedelta(months=1)
        while current_date.day != second_payment_day:
            current_date -= relativedelta(days=1)
        
        table_data.append([i + 1, current_date.strftime("%d.%m.%Y"), f"{payment_amount:.2f}"])

    current_date += relativedelta(months=1)
    while current_date.day != second_payment_day:
        current_date -= relativedelta(days=1)

    table_data.append([num_payments, current_date.strftime("%d.%m.%Y"), f"{last_payment:.2f}"])

    return table_data

def format_currency(amount):
    return f"{amount:,.2f}".replace(",", " ")



def generate_contract(data, template_path, output_path, payments):
    today = datetime.now().strftime("%d.%m.%Y")
    data["today"] = today
    discount = int(data.get("скидка", 0))

    if data.get('сумма бонус', '') == '':
        data['сумма бонус'] = '0'

    fio_parts = data.get("ФИО", "").split()
    if len(fio_parts) < 3:
        raise ValueError(f"ФИО должно содержать минимум 3 части (Фамилия Имя Отчество), получили: {data.get('ФИО')}")

    data["инициалы"] = f"{fio_parts[0]} {fio_parts[1][0]}. {fio_parts[2][0]}."

    data['сумма юристы'] = str(int(data['сумма юристы']) - discount)
    data["сумма"] = str(int(data["сумма бонус"]) + int(data["сумма юристы"]))
    data["words_sum"] = number_to_words(int(data['сумма юристы']))

    # --- Проверка шаблона ---
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)
    replace_text_in_paragraphs(doc, data)
    insert_table_after_heading(doc, payments)

    doc.save(output_path)
    logger.error(f"Document saved: {output_path}")

def apply_style_to_runs(paragraph):
    """Применяет стиль Montserrat 10pt ко всем runs в параграфе."""
    for run in paragraph.runs:
        apply_montserrat_to_run(run)


def set_run_font(run, font_name='Montserrat', font_size=10):
    """Устанавливает шрифт и размер для одного run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    rFonts = run._element.rPr.rFonts
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)


def replace_text_in_paragraphs(doc, data):
    """Заменяет плейсхолдеры вида {{ключ}} на значения из data во всём документе."""
    def process_paragraph(paragraph):
        full_text = paragraph.text
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if value is None:
                value = ""
            full_text = full_text.replace(placeholder, str(value))

        if full_text != paragraph.text:
            for _ in range(len(paragraph.runs)):
                paragraph.runs[0].clear()
                paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)

            run = paragraph.add_run(full_text)
            set_run_font(run)

    # Параграфы вне таблиц
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    # Параграфы внутри таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)


def apply_style_to_table_cells(doc):
    """Применяет шрифт Montserrat 10pt ко всем runs в таблицах документа."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        apply_montserrat_to_run(run)


def ensure_table_grid_style(doc):
    """Добавляет стандартный стиль Table Grid, если его нет в docx-шаблоне."""
    try:
        doc.styles["Table Grid"]
        return
    except KeyError:
        pass

    default_doc = Document()
    try:
        default_table_grid = default_doc.styles["Table Grid"]
    except KeyError:
        logger.warning("Default docx styles do not contain Table Grid")
        return

    doc.styles.element.append(deepcopy(default_table_grid.element))


def apply_table_grid_style(doc, table):
    """Применяет сетку таблицы как в прежней генерации договора."""
    ensure_table_grid_style(doc)
    try:
        table.style = "Table Grid"
    except KeyError:
        logger.warning("Table Grid style is missing in docx template; applying table borders manually")
    apply_table_borders(table)


def apply_table_borders(table):
    """Проставляет видимые границы таблицы на уровне таблицы и каждой ячейки."""
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_look = tbl_pr.first_child_found_in("w:tblLook")
    if tbl_look is None:
        tbl_look = OxmlElement("w:tblLook")
        tbl_pr.append(tbl_look)
    tbl_look.set(qn("w:firstRow"), "1")
    tbl_look.set(qn("w:lastRow"), "0")
    tbl_look.set(qn("w:firstColumn"), "0")
    tbl_look.set(qn("w:lastColumn"), "0")
    tbl_look.set(qn("w:noHBand"), "0")
    tbl_look.set(qn("w:noVBand"), "1")

    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)

    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        ensure_border(tbl_borders, border_name)

    apply_table_grid_columns(table)

    for row in table.rows:
        for col_idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_w)
            tc_w.set(qn("w:w"), str(get_payment_table_col_width(col_idx)))
            tc_w.set(qn("w:type"), "dxa")

            tc_borders = tc_pr.first_child_found_in("w:tcBorders")
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)

            for border_name in ("top", "left", "bottom", "right"):
                ensure_border(tc_borders, border_name)


def apply_table_grid_columns(table):
    if table._tbl.tr_lst:
        col_count = len(table._tbl.tr_lst[0].tc_lst)
    else:
        col_count = len(table.columns)

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, tbl_grid)

    for child in list(tbl_grid):
        tbl_grid.remove(child)

    for col_idx in range(col_count):
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(get_payment_table_col_width(col_idx)))
        tbl_grid.append(grid_col)


def get_payment_table_col_width(col_idx):
    return (1100, 3300, 2600)[col_idx] if col_idx < 3 else 2400


def ensure_border(parent, border_name):
    tag = f"w:{border_name}"
    border = parent.find(qn(tag))
    if border is None:
        border = OxmlElement(tag)
        parent.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "4")
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), "auto")


def insert_table_after_heading(doc, table_data):
    """Вставляет таблицу после заголовка `ГРАФИК ПЛАТЕЖЕЙ`."""
    heading_text = "ГРАФИК ПЛАТЕЖЕЙ"

    for paragraph in doc.paragraphs:
        if heading_text in paragraph.text:
            table = doc.add_table(rows=len(table_data) + 1, cols=3)
            apply_table_grid_style(doc, table)

            headers = ["П/П", "Дата платежа", "Сумма платежа"]
            for col_idx, header in enumerate(headers):
                cell = table.cell(0, col_idx)
                cell.text = header
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                cell.paragraphs[0].alignment = 1  # по центру

            for row_idx, row_data in enumerate(table_data, start=1):
                for col_idx, cell_data in enumerate(row_data):
                    cell = table.cell(row_idx, col_idx)
                    if cell_data is None:
                        cell_data = ""
                    if col_idx == 2:
                        try:
                            cell_data = str(int(float(cell_data)))
                        except:
                            cell_data = str(cell_data)
                    cell.text = str(cell_data)
                    cell.paragraphs[0].alignment = 1 

            apply_table_borders(table)
            paragraph._element.addnext(table._element)
            break


def replace_placeholders_in_runs(runs, data):
    """Заменяет плейсхолдеры в конкретных runs."""
    for run in runs:
        original_text = run.text
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if value is None:
                value = ""
            if placeholder in original_text:
                run.text = original_text.replace(placeholder, str(value))
                apply_montserrat_to_run(run)


def apply_montserrat_to_run(run):
    """Применяет шрифт Montserrat 10pt к одному run."""
    run.font.name = 'Montserrat'
    run.font.size = Pt(10)

    rFonts = run._element.rPr.rFonts
    rFonts.set(qn('w:ascii'), 'Montserrat')
    rFonts.set(qn('w:hAnsi'), 'Montserrat')
    rFonts.set(qn('w:cs'), 'Montserrat')
    rFonts.set(qn('w:eastAsia'), 'Montserrat')
    
    
    
def get_phone_number(contact_id):
    """Получает номер телефона по contact_id из Битрикс24"""
    contact_url = f"{WEBHOOK_URL}crm.contact.get.json?ID={contact_id}"

    response = requests.get(contact_url)
    
    if response.status_code == 200:
        try:
            contact_data = response.json().get('result', {})
            phone_list = contact_data.get("PHONE", [])
            if phone_list:
                return phone_list[0].get("VALUE", "")  
        except requests.exceptions.JSONDecodeError as e:
            return str(e)

    return ""



def format_date(date_str):
    if not date_str:
        return ""
    try:
        return datetime.fromisoformat(date_str).strftime("%d.%m.%Y")
    except ValueError:
        return date_str
    
    

def upload_to_bitrix(deal_id, file_path, field_id, payment_table):
    """Загрузка файла в сделку Bitrix через JSON + fileData"""

    if not os.path.exists(file_path):
        return {"status": "error", "message": f"Файл не найден: {file_path}"}

    with open(file_path, "rb") as f:
        encoded = base64.b64encode(f.read()).decode("utf-8")

    try:
        second_payment = payment_table[1][2]
    except Exception:
        second_payment = None

    fields = {
        field_id: {
            "fileData": [os.path.basename(file_path), encoded]
        }
    }

    if second_payment is not None:
        fields["UF_CRM_1745841297007"] = second_payment

    payload = {
        "id": deal_id,
        "fields": fields
    }

    url = f"{WEBHOOK_URL}crm.deal.update.json"
    response = requests.post(url, json=payload)

    try:
        result = response.json()
    except:
        return {
            "status": "error",
            "message": "Bitrix вернул не-JSON",
            "raw": response.text
        }

    if result.get("error"):
        return {
            "status": "error",
            "message": result.get("error_description", "Ошибка Bitrix"),
            "details": result
        }

    return {"status": "success", "message": "Файл загружен"}


def number_to_words(num):
    if num == 0:
        return "ноль"

    units = (
        "", "один", "два", "три", "четыре", "пять",
        "шесть", "семь", "восемь", "девять"
    )
    teens = (
        "десять", "одиннадцать", "двенадцать", "тринадцать",
        "четырнадцать", "пятнадцать", "шестнадцать",
        "семнадцать", "восемнадцать", "девятнадцать"
    )
    tens = (
        "", "", "двадцать", "тридцать", "сорок",
        "пятьдесят", "шестьдесят", "семьдесят",
        "восемьдесят", "девяносто"
    )
    hundreds = (
        "", "сто", "двести", "триста", "четыреста",
        "пятьсот", "шестьсот", "семьсот", "восемьсот",
        "девятьсот"
    )
    thousands_forms = ("тысяча", "тысячи", "тысяч")
    millions_forms = ("миллион", "миллиона", "миллионов")

    def get_form(number, forms):
        if 11 <= number % 100 <= 19:
            return forms[2]
        elif number % 10 == 1:
            return forms[0]
        elif 2 <= number % 10 <= 4:
            return forms[1]
        else:
            return forms[2]

    def three_digit_number_to_words(n):
        result = []
        if n >= 100:
            result.append(hundreds[n // 100])
            n %= 100
        if 10 <= n < 20:
            result.append(teens[n - 10])
        else:
            if n >= 20:
                result.append(tens[n // 10])
            if n % 10 > 0:
                result.append(units[n % 10])
        
        
        return " ".join(result).strip()


    result = []

    if num >= 1_000_000:
        millions = num // 1_000_000
        result.append(f"{three_digit_number_to_words(millions)} {get_form(millions, millions_forms)}")
        num %= 1_000_000

    if num >= 1_000:
        thousands = num // 1_000
        thousands_text = three_digit_number_to_words(thousands)

        words = thousands_text.split()
        for i in range(len(words)):
            if words[i] == "один":
                words[i] = "одна"
            elif words[i] == "два":
                words[i] = "две"
        thousands_text = " ".join(words)

        result.append(f"{thousands_text} {get_form(thousands, thousands_forms)}")
        num %= 1_000

    if num > 0:
        result.append(three_digit_number_to_words(num))

    return " ".join(result).strip()
    
    
    
    
    
def parse_creditors_and_calculate(request):
    creditor_names = request.POST.getlist("creditors[]")
    creditor_amounts = request.POST.getlist("amounts[]")
    creditor_dates = request.POST.getlist("court_date[]")

    # Указанная вручную сумма распределения
    try:
        distribution_sum = float(request.POST.get("distribution_sum", 0))
    except ValueError:
        distribution_sum = 0

    creditors_data = []
    total_debt = 0

    for name, amount, date in zip(creditor_names, creditor_amounts, creditor_dates):
        if not name.strip():
            continue

        try:
            debt = float(amount)
        except ValueError:
            debt = 0

        creditors_data.append({
            "name": name.strip(),
            "debt": debt,
            "date": date  # 👈 добавили дату
        })

        total_debt += debt

    # Расчёт пропорциональной выплаты
    for c in creditors_data:
        if total_debt > 0:
            proportion = c["debt"] / total_debt
        else:
            proportion = 0

        c["pay"] = round(distribution_sum * proportion, 2)

    return creditors_data, total_debt, distribution_sum 



#LEGENDA--------------------------------------------------------------------------------------------------------------------------------

OPENAI_API_KEY = "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"  # лучше перенести в переменные окружения
BITRIX_WEBHOOK_URL = "https://prav-buro.bitrix24.ru/rest/24/pa1x5irnfpbcnh27/"

TOR_PROXIES = {
    'http': 'socks5h://127.0.0.1:9050',
    'https': 'socks5h://127.0.0.1:9050'
    }


def get_deal_data_from_bitrix(post_data):
    """
    Извлекает ID сделки из POST-данных Bitrix24 и возвращает данные сделки
    """
    document_id_2 = post_data.get('document_id[2]')
    if not document_id_2:
        return None, 'document_id[2] not found'

    deal_id_match = re.search(r'(?:DEAL[_-])?(\d+)', str(document_id_2))
    if not deal_id_match:
        return None, 'Invalid deal ID format'

    deal_id = deal_id_match.group(1)

    webhook_url = f"{BITRIX_WEBHOOK_URL}crm.deal.get.json?ID={deal_id}"
    response = requests.get(webhook_url)

    if response.status_code != 200:
        return None, f"Bitrix24 request failed with status {response.status_code}"

    try:
        deal_data = response.json().get('result', {})
        return deal_data, None
    except json.JSONDecodeError:
        return None, 'Invalid JSON response from Bitrix'

@csrf_exempt
@require_POST
def parse_legenda(request):
    print(">>> parse_legenda HIT", request.method)
    try:
        post_data = request.POST.dict()
        deal_data, error = get_deal_data_from_bitrix(post_data)
        if error:
            return JsonResponse({"error": error}, status=400)

        legenda = deal_data.get("COMMENTS")
        deal_id = deal_data.get("ID")

        if not legenda:
            return JsonResponse({"error": "В сделке нет текста легенды (COMMENTS)"}, status=400)
        if not deal_id:
            return JsonResponse({"error": "Не найден ID сделки"}, status=400)

        # === Шаг 2: System prompt ===
        system_prompt = (
            "Ты — аналитик юридической компании. "
            "Твоя задача — строго по тексту легенды извлечь данные для CRM в формате JSON. "
            "Все данные должны быть напрямую подтверждены текстом, а не основаны на догадках.\n\n"
            "📌 Правила извлечения:\n"
            "1. Читай весь текст внимательно и используй только ту информацию, которая прямо указана.\n"
            "2. Если есть сомнения или нет точного упоминания — заполняй поле указанным fallback-значением.\n"
            "3. Для чисел бери только те, что явно относятся к данному полю, игнорируй все остальные.\n"
            "4. Разрешено перефразировать ответы для улучшения читаемости в полях:\n"
            "   UF_CRM_1754647579070, UF_CRM_1754647590990, UF_CRM_1754647601622,\n"
            "   UF_CRM_1754647621350, UF_CRM_1754647681566, UF_CRM_1754647691574.\n"
            "5. Для остальных полей текст должен быть максимально близок к оригиналу.\n"
            "6. Ответ всегда в виде одного корректного JSON-объекта без пояснений и комментариев.\n"
            "7. Верни объект СТРОГО с этими ключами (все 11 ключей обязательны). Если данных нет — используй fallback.\n\n"
            "📌 Формат и fallback-значения:\n"
            "- UF_CRM_1754647579070: боль клиента. Если нет — 'не обнаружено'\n"
            "- UF_CRM_1754647590990: краткая характеристика клиента, как он общается - если нет данных — 'адекватный'\n"
            "- UF_CRM_1754647601622: имущество клиента. Если нет — 'не обнаружено'\n"
            "- UF_CRM_1754647621350: состоит ли клиент в браке и есть ли совместное имущество. Если нет данных — 'неизвестно'\n"
            "- UF_CRM_1754647636597: доход клиента. Если нет данных или не уверен — 'неизвестно'\n"
            "- UF_CRM_1754647649551: знает ли клиент про блокировку карт в банкротстве. Если нет — 'в диалоге не упоминалось'\n"
            "- UF_CRM_1754647663541: сделки за 3 года. Если нет — 'не обнаружено'\n"
            "- UF_CRM_1754647671862: дети на иждивении. Если нет — 'не обнаружено'\n"
            "- UF_CRM_1754647681566: что обещано клиенту. Если нет — 'просто пройти процедуру'\n"
            "- UF_CRM_1754647691574: пожелания клиента. Если нет — 'пожеланий нет'\n"
            "- UF_CRM_1754647902223: сделки супруга за 3 года. Если не в браке — 'не в браке'\n"
        )

        # === Нормализация: гарантируем полный набор ключей ===
        expected_defaults = {
            "UF_CRM_1754647579070": "не обнаружено",
            "UF_CRM_1754647590990": "адекватный",
            "UF_CRM_1754647601622": "не обнаружено",
            "UF_CRM_1754647621350": "неизвестно",
            "UF_CRM_1754647636597": "неизвестно",
            "UF_CRM_1754647649551": "в диалоге не упоминалось",
            "UF_CRM_1754647663541": "не обнаружено",
            "UF_CRM_1754647671862": "не обнаружено",
            "UF_CRM_1754647681566": "просто пройти процедуру",
            "UF_CRM_1754647691574": "пожеланий нет",
            "UF_CRM_1754647902223": "не в браке",
        }

        parsed_fields = {}

        # === Шаг 3: Запрос к OpenAI ===
        try:
            payload = {
            "model": "gpt-4o",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": legenda},
            ],
            "temperature": 0.2,
        }

            openai_kwargs = {
                "headers": {
                    "Authorization": f"Bearer {OPENAI_API_KEY}",
                    "Content-Type": "application/json",
                },
                "json": payload,
                "timeout": 60,
            }

            if USE_PROXIES:
                openai_kwargs["proxies"] = TOR_PROXIES

            openai_response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                **openai_kwargs
            )

            openai_response.raise_for_status()

            raw_output = openai_response.json()["choices"][0]["message"]["content"].strip()
            cleaned_output = re.sub(r"^```json\s*|\s*```$", "", raw_output, flags=re.IGNORECASE).strip()

            parsed_fields = json.loads(cleaned_output)
            if not isinstance(parsed_fields, dict):
                raise ValueError("OpenAI вернул не JSON-объект")

        except Exception as e:
            print("OpenAI error:", e)
            parsed_fields = {}

        # Добиваем отсутствующие ключи fallback-ами
        normalized_fields = dict(expected_defaults)
        for k, v in (parsed_fields or {}).items():
            if k in expected_defaults:
                normalized_fields[k] = v if v is not None and str(v).strip() != "" else expected_defaults[k]

        # === Шаг 4: Обновляем сделку в Bitrix ===
        bitrix_update_ok = False
        try:
            bitrix_payload = {"id": deal_id}
            for field_code, field_value in normalized_fields.items():
                bitrix_payload[f"fields[{field_code}]"] = field_value

            bitrix_update_url = f"{BITRIX_WEBHOOK_URL}crm.deal.update.json"
            bitrix_response = requests.post(bitrix_update_url, data=bitrix_payload)
            # ВАЖНО: битрикс может вернуть 200, но с "error" внутри
            try:
                bitrix_json = bitrix_response.json()
            except Exception:
                bitrix_json = {}

            print("Bitrix update resp:", bitrix_response.text)

            bitrix_response.raise_for_status()
            if "error" in bitrix_json:
                raise Exception(f"{bitrix_json.get('error')}: {bitrix_json.get('error_description')}")

            bitrix_update_ok = True

        except Exception as e:
            print("Bitrix update error:", e)
            bitrix_update_ok = False

        # === Шаг 4.1: Ждём, пока обновления реально станут видны (защита от гонки) ===
        # Если не нужно — можешь убрать этот блок, но он помогает, когда DG берет старые данные.
        waited_ok = False
        try:
            if bitrix_update_ok:
                deal_get_url = f"{BITRIX_WEBHOOK_URL}crm.deal.get.json"
                for _ in range(6):  # 6 попыток
                    get_resp = requests.post(deal_get_url, data={"id": deal_id})
                    get_json = get_resp.json()
                    current = get_json.get("result") or {}

                    # Проверяем, что все ключи совпали со значениями, которые записывали
                    if all(str(current.get(k, "")) == str(normalized_fields.get(k, "")) for k in expected_defaults.keys()):
                        waited_ok = True
                        break

                    time.sleep(0.5)
        except Exception as e:
            print("Bitrix get/wait error:", e)

        # === Шаг 5: Генерация документа ===
        # КЛЮЧЕВОЕ: прокидываем values прямо в документ, чтобы DG не зависел от того,
        # успела ли сущность обновиться/проиндексироваться.
        template_id = 40
        entity_type_id = 2  # Сделка
        document_url = f"{BITRIX_WEBHOOK_URL}crm.documentgenerator.document.add.json"

        download_url = None
        pdf_url = None

        try:
            document_payload = {
                "templateId": template_id,
                "entityTypeId": entity_type_id,
                "entityId": deal_id,
                "values": {
                    "SAVE_IN_ENTITY": "Y",
                    **normalized_fields,   # <-- важная часть
                },
            }

            doc_response = requests.post(document_url, json=document_payload)
            print("Doc gen resp:", doc_response.text)

            doc_response.raise_for_status()
            doc_json = doc_response.json()
            if "error" in doc_json:
                raise Exception(f"{doc_json.get('error')}: {doc_json.get('error_description')}")

            doc_result = doc_json.get("result")
            if doc_result and "document" in doc_result:
                doc_data = doc_result["document"]
                download_url = doc_data.get("downloadUrl")
                pdf_url = doc_data.get("pdfUrl")

        except Exception as e:
            return JsonResponse({"error": "Не удалось сгенерировать документ", "detail": str(e)}, status=502)

        return JsonResponse(
            {
                "success": True,
                "data_written": normalized_fields,
                "bitrix_update_ok": bitrix_update_ok,
                "bitrix_waited_ok": waited_ok,
                "document": {
                    "downloadUrl": download_url,
                    "pdfUrl": pdf_url,
                },
            }
        )

    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)
