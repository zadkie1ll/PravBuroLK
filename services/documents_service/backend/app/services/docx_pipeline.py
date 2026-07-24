from __future__ import annotations

import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Any

from dateutil.relativedelta import relativedelta
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

logger = logging.getLogger(__name__)


def calculate_payments(
    num_payments: int,
    total_amount: float,
    discount: float,
    start_date: str,
    first_payment: float,
    second_payment_day: int | str,
) -> list[list[Any]]:
    if num_payments == 1 and first_payment >= (total_amount - discount):
        return [[1, start_date, f"{first_payment:.2f}"]]

    remaining_amount = (total_amount - discount) - first_payment
    remaining_payments = num_payments - 1

    if remaining_payments == 0:
        return [[1, start_date, f"{first_payment:.2f}"]]

    payment_amount = round(remaining_amount / remaining_payments, -2)

    adjusted_total = first_payment + payment_amount * (remaining_payments - 1)
    last_payment = (total_amount - discount) - adjusted_total

    first_date = datetime.strptime(start_date, "%d.%m.%Y").date()
    table_data: list[list[Any]] = [[1, first_date.strftime("%d.%m.%Y"), f"{first_payment:.2f}"]]

    second_payment_day = int(second_payment_day)
    second_date = (first_date + relativedelta(months=1)).replace(day=second_payment_day)

    while second_date.day != second_payment_day:
        second_date -= relativedelta(days=1)

    table_data.append([2, second_date.strftime("%d.%m.%Y"), f"{payment_amount:.2f}"])

    current_date = second_date
    for i in range(2, num_payments - 1):
        current_date += relativedelta(months=1)
        while current_date.day != second_payment_day:
            current_date -= relativedelta(days=1)
        table_data.append([i + 1, current_date.strftime("%d.%m.%Y"), f"{payment_amount:.2f}"])

    current_date += relativedelta(months=1)
    while current_date.day != second_payment_day:
        current_date -= relativedelta(days=1)

    table_data.append([num_payments, current_date.strftime("%d.%m.%Y"), f"{last_payment:.2f}"])

    return table_data


def format_date(date_str: str | None) -> str:
    if not date_str:
        return ""
    try:
        return datetime.fromisoformat(date_str).strftime("%d.%m.%Y")
    except ValueError:
        return date_str


def number_to_words(num: int) -> str:
    if num == 0:
        return "ноль"

    units = ("", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять")
    teens = (
        "десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать",
        "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать",
    )
    tens = (
        "", "", "двадцать", "тридцать", "сорок", "пятьдесят", "шестьдесят",
        "семьдесят", "восемьдесят", "девяносто",
    )
    hundreds = (
        "", "сто", "двести", "триста", "четыреста", "пятьсот", "шестьсот",
        "семьсот", "восемьсот", "девятьсот",
    )
    thousands_forms = ("тысяча", "тысячи", "тысяч")
    millions_forms = ("миллион", "миллиона", "миллионов")

    def get_form(number: int, forms: tuple[str, str, str]) -> str:
        if 11 <= number % 100 <= 19:
            return forms[2]
        elif number % 10 == 1:
            return forms[0]
        elif 2 <= number % 10 <= 4:
            return forms[1]
        return forms[2]

    def three_digit_number_to_words(n: int) -> str:
        result = []
        if n >= 100:
            result.append(hundreds[n // 100])
            n %= 100
        if 10 <= n < 20:
            result.append(teens[n - 10])
        else:
            if n >= 20:
                result.append(tens[n // 10])
            if n % 10 > 0:
                result.append(units[n % 10])
        return " ".join(result).strip()

    result = []

    if num >= 1_000_000:
        millions = num // 1_000_000
        result.append(f"{three_digit_number_to_words(millions)} {get_form(millions, millions_forms)}")
        num %= 1_000_000

    if num >= 1_000:
        thousands = num // 1_000
        thousands_text = three_digit_number_to_words(thousands)
        words = thousands_text.split()
        for i in range(len(words)):
            if words[i] == "один":
                words[i] = "одна"
            elif words[i] == "два":
                words[i] = "две"
        thousands_text = " ".join(words)
        result.append(f"{thousands_text} {get_form(thousands, thousands_forms)}")
        num %= 1_000

    if num > 0:
        result.append(three_digit_number_to_words(num))

    return " ".join(result).strip()


def _set_run_font(run, font_name: str = "Montserrat", font_size: int = 10) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r_fonts = run._element.rPr.rFonts
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:cs"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)


def _replace_text_in_paragraphs(doc: Document, data: dict[str, Any]) -> None:
    def process_paragraph(paragraph) -> None:
        full_text = paragraph.text
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if value is None:
                value = ""
            full_text = full_text.replace(placeholder, str(value))

        if full_text != paragraph.text:
            for _ in range(len(paragraph.runs)):
                paragraph.runs[0].clear()
                paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
            run = paragraph.add_run(full_text)
            _set_run_font(run)

    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)


def _insert_table_after_heading(doc: Document, table_data: list[list[Any]]) -> None:
    heading_text = "ГРАФИК ПЛАТЕЖЕЙ"

    for paragraph in doc.paragraphs:
        if heading_text in paragraph.text:
            table = doc.add_table(rows=len(table_data) + 1, cols=3)
            table.style = "Table Grid"

            headers = ["П/П", "Дата платежа", "Сумма платежа"]
            for col_idx, header in enumerate(headers):
                cell = table.cell(0, col_idx)
                cell.text = header
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                cell.paragraphs[0].alignment = 1

            for row_idx, row_data in enumerate(table_data, start=1):
                for col_idx, cell_data in enumerate(row_data):
                    cell = table.cell(row_idx, col_idx)
                    if cell_data is None:
                        cell_data = ""
                    if col_idx == 2:
                        try:
                            cell_data = str(int(float(cell_data)))
                        except (TypeError, ValueError):
                            cell_data = str(cell_data)
                    cell.text = str(cell_data)
                    cell.paragraphs[0].alignment = 1

            paragraph._element.addnext(table._element)
            break


def generate_contract(
    data: dict[str, Any],
    template_path: str,
    output_path: str,
    payments: list[list[Any]],
) -> None:
    today = datetime.now().strftime("%d.%m.%Y")
    data["today"] = today
    discount = int(data.get("скидка", 0))

    if data.get("сумма бонус", "") == "":
        data["сумма бонус"] = "0"

    fio_parts = data.get("ФИО", "").split()
    if len(fio_parts) < 3:
        raise ValueError(f"ФИО должно содержать минимум 3 части (Фамилия Имя Отчество), получили: {data.get('ФИО')}")

    data["инициалы"] = f"{fio_parts[0]} {fio_parts[1][0]}. {fio_parts[2][0]}."
    data["сумма юристы"] = str(int(data["сумма юристы"]) - discount)
    data["сумма"] = str(int(data["сумма бонус"]) + int(data["сумма юристы"]))
    data["words_sum"] = number_to_words(int(data["сумма юристы"]))

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)
    _replace_text_in_paragraphs(doc, data)
    _insert_table_after_heading(doc, payments)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info("Contract document saved: %s", output_path)
